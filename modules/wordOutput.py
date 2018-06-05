import datetime
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from docx import Document
from docx.shared import Pt, Inches
import os


filetime = datetime.datetime.now().strftime("_%m_%d_%y_%H_%M_%S")
fileout = os.getcwd() + '/report_Output/sslscan_Report' + filetime + '.docx'

def remove_row(table, row):
    tbl = table._tbl
    tr = row._tr
    tbl.remove(tr)


def set_column_width(column, width):
    column.width = width
    for cell in column.cells:
        cell.width = width


def create_document(scaninput, reportout=fileout):
    length1 = int(len(scaninput) / 3)

    shading2 = parse_xml(r'<w:shd {} w:fill="00000a"/>'.format(nsdecls('w')))
    document = Document()

    table = document.add_table(rows=length1 + 1, cols=3)
    table.style = 'Table Grid'

    set_column_width(table.columns[0], Inches(1.2))
    set_column_width(table.columns[1], Inches(5))
    set_column_width(table.columns[2], Inches(1.2))

    cells1 = table.rows[0].cells
    cells1[0].text = 'Host:'
    cells1[1].text = 'Weak Protocols & Ciphers:'
    cells1[2].text = 'Certificate Information:'
    for x in range(0, 3):
        cell2 = table.cell(0, x)
        run = cell2.paragraphs[0].runs[0]
        run.font.name = 'Hind'
        run.font.size = Pt(11)
        run.font.bold = True
        shading1 = 'shading' + str(x)
        shading1 = parse_xml(r'<w:shd {} w:fill="f2f2f2"/>'.format(nsdecls('w')))
        table.rows[0].cells[x]._tc.get_or_add_tcPr().append(shading1)


    numrows = 1
    for counter in range(1, length1+1):
        cells1 = table.rows[numrows].cells

        test1 = scaninput.get('ciphers'+str(counter))
        test2 = test1.split(' •')
        weak1 = []
        for x in test2:
            if 'SSLv2' in x or 'SSLv3' in x or 'TLSv1.0' in x or 'RC4' in x or 'CBC' in x:
                weak1.append('•' + x)
        weak2 = '\n'.join(weak1)

        cells1[0].text = scaninput.get('host' + str(counter))
        cells1[1].text = weak2
        cells1[2].text = scaninput.get('cert' + str(counter))

        if not cells1[1].text:
            remove_row(table, table.rows[numrows])
            numrows -= 1
        numrows += 1

    for x in range(0, 3):
        for y in range(1, numrows):
            cell2 = table.cell(y, x)
            run = cell2.paragraphs[0].runs[0]
            run.font.name = 'Hind'
            run.font.size = Pt(10)
            mod = y % 2
            if mod > 0:
                shading2 = 'shading' + str(x)
                shading2 = parse_xml(r'<w:shd {} w:fill="ffffff"/>'.format(nsdecls('w')))
                table.rows[y].cells[x]._tc.get_or_add_tcPr().append(shading2)
            else:
                shading3 = 'shading' + str(x)
                shading3 = parse_xml(r'<w:shd {} w:fill="f2f2f2"/>'.format(nsdecls('w')))
                table.rows[y].cells[x]._tc.get_or_add_tcPr().append(shading3)

    document.save(reportout)
    return reportout
